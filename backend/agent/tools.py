import os
from pypdf import PdfReader
import docx
from backend.agent.embedder import embed_document, is_document_embedded
from backend.agent.retriever import search_documents as _search_documents
from openai import OpenAI


def list_documents(folder_path: str):
    home = os.path.expanduser("~")

    path_map = {
        "downloads": os.path.join(home, "Downloads"),
        "desktop": os.path.join(home, "Desktop"),
        "documents": os.path.join(home, "Documents")
    }

    if folder_path.lower() == "all":
        folders_to_check = path_map.values()
    else:
        resolved = path_map.get(folder_path.lower(), None)
        if not resolved:
            raise Exception(f"Invalid folder: {folder_path}. Must be downloads, desktop, documents, or all.")
        folders_to_check = [resolved]

    files = []
    for folder in folders_to_check:
        if os.path.exists(folder):
            for f in os.listdir(folder):
                if f.lower().endswith((".pdf", ".docx")):
                    files.append({
                        "name": f,
                        "path": os.path.join(folder, f)
                    })

    return files


def read_pdf(path: str):
    if not os.path.exists(path):
        raise Exception(f"File not found: {path}")

    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""

    return {
        "path": path,
        "content": text[:8000]
    }


def read_text_file(path: str):
    if not os.path.exists(path):
        raise Exception(f"File not found: {path}")

    with open(path, "r", encoding="utf-8") as f:
        return {
            "path": path,
            "content": f.read()
        }


def read_word_file(file_path: str):
    if not os.path.exists(file_path):
        raise Exception(f"File not found: {file_path}")
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return {
            "path": file_path,
            "content": "\n".join(full_text)[:8000]
        }
    except Exception as e:
        return f"Error parsing Word file: {str(e)}"


def embed_and_index(file_path: str):
    if is_document_embedded(file_path):
        return {"status": "already_indexed", "file_path": file_path}

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        result = read_pdf(file_path)
    elif ext == ".docx":
        result = read_word_file(file_path)
    elif ext == ".txt":
        result = read_text_file(file_path)
    else:
        raise Exception(f"Unsupported file type: {ext}")

    content = result.get("content", "") if isinstance(result, dict) else str(result)
    chunks_created = embed_document(file_path, content)

    return {
        "status": "indexed",
        "file_path": file_path,
        "chunks_created": chunks_created
    }


def search_documents(query: str):
    hits = _search_documents(query, n_results=5)

    if not hits:
        return {"status": "no_results", "message": "No relevant content found in indexed documents."}

    return {
        "status": "success",
        "results": hits
    }

def read_image(file_path: str):
    import base64
    if not os.path.exists(file_path):
        raise Exception(f"File not found: {file_path}")

    with open(file_path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    ext = os.path.splitext(file_path)[1].lower().replace(".", "")
    if ext == "jpg":
        ext = "jpeg"

    client = OpenAI()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/{ext};base64,{image_data}"
                        }
                    },
                    {
                        "type": "text",
                        "text": "Describe this image in detail. Extract any text, data, charts, or important information visible."
                    }
                ]
            }
        ],
        max_tokens=1000
    )

    return {
        "path": file_path,
        "content": response.choices[0].message.content
    }